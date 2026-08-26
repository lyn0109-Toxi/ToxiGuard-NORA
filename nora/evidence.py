from __future__ import annotations

import csv
import hashlib
import io
import json
import re
import uuid
from dataclasses import asdict, dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader


@dataclass
class SourceSegment:
    segment_id: str
    location: str
    text: str

    def to_dict(self) -> dict[str, str]:
        return asdict(self)

    @classmethod
    def from_dict(cls, payload: dict[str, Any]) -> "SourceSegment":
        return cls(
            segment_id=str(payload.get("segment_id", "")),
            location=str(payload.get("location", "")),
            text=str(payload.get("text", "")),
        )


@dataclass
class DocumentRecord:
    document_id: str
    name: str
    media_type: str
    extension: str
    sha256: str
    byte_size: int
    extracted_text: str
    segments: list[SourceSegment] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    uploaded_at_utc: str = ""

    def to_dict(self) -> dict[str, Any]:
        payload = asdict(self)
        payload["segments"] = [segment.to_dict() for segment in self.segments]
        return payload

    @classmethod
    def from_dict(cls, payload: dict[str, Any]) -> "DocumentRecord":
        return cls(
            document_id=str(payload.get("document_id", "")),
            name=str(payload.get("name", "")),
            media_type=str(payload.get("media_type", "application/octet-stream")),
            extension=str(payload.get("extension", "")),
            sha256=str(payload.get("sha256", "")),
            byte_size=int(payload.get("byte_size", 0) or 0),
            extracted_text=str(payload.get("extracted_text", "")),
            segments=[SourceSegment.from_dict(item) for item in payload.get("segments", [])],
            warnings=[str(item) for item in payload.get("warnings", [])],
            uploaded_at_utc=str(payload.get("uploaded_at_utc", "")),
        )


@dataclass
class EvidenceAssertion:
    assertion_id: str
    category: str
    field_path: str
    label_ko: str
    proposed_value: str
    value_type: str
    source_document_id: str
    source_document_name: str
    source_location: str
    source_excerpt: str
    confidence: float
    review_status: str = "제안됨"
    reviewer_note: str = ""

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)

    @classmethod
    def from_dict(cls, payload: dict[str, Any]) -> "EvidenceAssertion":
        return cls(
            assertion_id=str(payload.get("assertion_id", "")),
            category=str(payload.get("category", "")),
            field_path=str(payload.get("field_path", "")),
            label_ko=str(payload.get("label_ko", "")),
            proposed_value=str(payload.get("proposed_value", "")),
            value_type=str(payload.get("value_type", "str")),
            source_document_id=str(payload.get("source_document_id", "")),
            source_document_name=str(payload.get("source_document_name", "")),
            source_location=str(payload.get("source_location", "")),
            source_excerpt=str(payload.get("source_excerpt", "")),
            confidence=float(payload.get("confidence", 0.0) or 0.0),
            review_status=str(payload.get("review_status", "제안됨")),
            reviewer_note=str(payload.get("reviewer_note", "")),
        )


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".csv", ".xlsx", ".xlsm", ".json"}


def _utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _new_id(prefix: str) -> str:
    return f"{prefix}-{uuid.uuid4().hex[:12]}"


def _normalize(text: str) -> str:
    text = (text or "").replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _chunk_lines(text: str, *, label_prefix: str = "텍스트", lines_per_chunk: int = 45) -> list[SourceSegment]:
    lines = [line.rstrip() for line in (text or "").splitlines()]
    segments: list[SourceSegment] = []
    for start in range(0, len(lines), lines_per_chunk):
        chunk = _normalize("\n".join(lines[start : start + lines_per_chunk]))
        if not chunk:
            continue
        end = min(len(lines), start + lines_per_chunk)
        segments.append(
            SourceSegment(
                segment_id=_new_id("SEG"),
                location=f"{label_prefix} {start + 1}-{end}행",
                text=chunk,
            )
        )
    if not segments and _normalize(text):
        segments.append(SourceSegment(_new_id("SEG"), label_prefix, _normalize(text)))
    return segments


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp949", "euc-kr", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _extract_pdf(data: bytes) -> tuple[list[SourceSegment], list[str]]:
    warnings: list[str] = []
    segments: list[SourceSegment] = []
    reader = PdfReader(io.BytesIO(data))
    if getattr(reader, "is_encrypted", False):
        try:
            reader.decrypt("")
        except Exception as exc:  # pragma: no cover - depends on encryption
            return [], [f"암호화된 PDF를 읽지 못했습니다: {exc}"]
    for index, page in enumerate(reader.pages, start=1):
        try:
            text = _normalize(page.extract_text() or "")
        except Exception as exc:  # pragma: no cover - malformed PDF edge case
            warnings.append(f"{index}페이지 텍스트 추출 실패: {exc}")
            text = ""
        if text:
            segments.append(SourceSegment(_new_id("SEG"), f"PDF {index}페이지", text))
        else:
            warnings.append(f"PDF {index}페이지에서 검색 가능한 텍스트를 찾지 못했습니다. 스캔 문서일 수 있습니다.")
    return segments, warnings


def _extract_docx(data: bytes) -> tuple[list[SourceSegment], list[str]]:
    document = Document(io.BytesIO(data))
    segments: list[SourceSegment] = []
    paragraph_buffer: list[str] = []
    paragraph_start = 1

    def flush_paragraphs(end_index: int) -> None:
        nonlocal paragraph_buffer, paragraph_start
        text = _normalize("\n".join(paragraph_buffer))
        if text:
            segments.append(
                SourceSegment(
                    _new_id("SEG"),
                    f"DOCX 문단 {paragraph_start}-{end_index}",
                    text,
                )
            )
        paragraph_buffer = []
        paragraph_start = end_index + 1

    for index, paragraph in enumerate(document.paragraphs, start=1):
        paragraph_buffer.append(paragraph.text)
        if len(paragraph_buffer) >= 25 or sum(len(item) for item in paragraph_buffer) >= 4000:
            flush_paragraphs(index)
    if paragraph_buffer:
        flush_paragraphs(len(document.paragraphs))

    for table_index, table in enumerate(document.tables, start=1):
        rows: list[str] = []
        for row_index, row in enumerate(table.rows, start=1):
            values = [re.sub(r"\s+", " ", cell.text).strip() for cell in row.cells]
            rows.append(" | ".join(values))
            if len(rows) >= 35:
                segments.append(
                    SourceSegment(
                        _new_id("SEG"),
                        f"DOCX 표 {table_index} 행 {row_index - len(rows) + 1}-{row_index}",
                        _normalize("\n".join(rows)),
                    )
                )
                rows = []
        if rows:
            segments.append(
                SourceSegment(
                    _new_id("SEG"),
                    f"DOCX 표 {table_index} 마지막 {len(rows)}행",
                    _normalize("\n".join(rows)),
                )
            )
    return segments, []


def _extract_csv(data: bytes) -> tuple[list[SourceSegment], list[str]]:
    text = _decode_text(data)
    reader = csv.reader(io.StringIO(text))
    rows = list(reader)
    segments: list[SourceSegment] = []
    for start in range(0, len(rows), 45):
        chunk = rows[start : start + 45]
        serialized = "\n".join(" | ".join(str(cell) for cell in row) for row in chunk)
        serialized = _normalize(serialized)
        if serialized:
            segments.append(
                SourceSegment(
                    _new_id("SEG"),
                    f"CSV 행 {start + 1}-{start + len(chunk)}",
                    serialized,
                )
            )
    return segments, []


def _extract_xlsx(data: bytes) -> tuple[list[SourceSegment], list[str]]:
    workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=False)
    segments: list[SourceSegment] = []
    warnings: list[str] = []
    for sheet in workbook.worksheets:
        buffer: list[str] = []
        start_row = 1
        for row_index, row in enumerate(sheet.iter_rows(values_only=False), start=1):
            values: list[str] = []
            for cell in row:
                value = cell.value
                if value is None:
                    values.append("")
                elif isinstance(value, str) and value.startswith("="):
                    values.append(f"{cell.coordinate}:{value}")
                else:
                    values.append(f"{cell.coordinate}:{value}")
            buffer.append(" | ".join(values))
            if len(buffer) >= 35:
                text = _normalize("\n".join(buffer))
                if text:
                    segments.append(
                        SourceSegment(
                            _new_id("SEG"),
                            f"Excel '{sheet.title}' {start_row}-{row_index}행",
                            text,
                        )
                    )
                buffer = []
                start_row = row_index + 1
        if buffer:
            text = _normalize("\n".join(buffer))
            if text:
                segments.append(
                    SourceSegment(
                        _new_id("SEG"),
                        f"Excel '{sheet.title}' {start_row}-{sheet.max_row}행",
                        text,
                    )
                )
        if sheet.max_row == 1 and sheet.max_column == 1 and sheet["A1"].value is None:
            warnings.append(f"Excel 시트 '{sheet.title}'가 비어 있습니다.")
    return segments, warnings


def _extract_json(data: bytes) -> tuple[list[SourceSegment], list[str]]:
    text = _decode_text(data)
    warnings: list[str] = []
    try:
        payload = json.loads(text)
        text = json.dumps(payload, ensure_ascii=False, indent=2)
    except json.JSONDecodeError as exc:
        warnings.append(f"JSON 파싱 실패; 일반 텍스트로 처리했습니다: {exc}")
    return _chunk_lines(text, label_prefix="JSON", lines_per_chunk=60), warnings


def extract_document(data: bytes, name: str, media_type: str | None = None) -> DocumentRecord:
    """Extract searchable text while preserving page/sheet/row-level provenance."""
    clean_name = Path(name or "document").name[:180]
    extension = Path(clean_name).suffix.lower()
    media_type = media_type or "application/octet-stream"
    digest = hashlib.sha256(data).hexdigest()
    warnings: list[str] = []

    try:
        if extension == ".pdf":
            segments, warnings = _extract_pdf(data)
        elif extension == ".docx":
            segments, warnings = _extract_docx(data)
        elif extension in {".xlsx", ".xlsm"}:
            segments, warnings = _extract_xlsx(data)
        elif extension == ".csv":
            segments, warnings = _extract_csv(data)
        elif extension == ".json":
            segments, warnings = _extract_json(data)
        elif extension in {".txt", ".md"}:
            segments = _chunk_lines(_decode_text(data), label_prefix=extension[1:].upper(), lines_per_chunk=50)
        else:
            text = _decode_text(data)
            segments = _chunk_lines(text, label_prefix="일반 텍스트", lines_per_chunk=50)
            warnings.append(f"공식 지원 형식이 아닌 '{extension or '확장자 없음'}' 파일을 일반 텍스트로 처리했습니다.")
    except Exception as exc:
        segments = []
        warnings.append(f"문서 처리 중 오류가 발생했습니다: {type(exc).__name__}: {exc}")

    extracted_text = "\n\n".join(
        f"[{segment.location}]\n{segment.text}" for segment in segments if segment.text
    )
    if not extracted_text:
        warnings.append("검색 가능한 텍스트가 없어 근거 Assertion을 자동 추출할 수 없습니다.")

    return DocumentRecord(
        document_id=_new_id("DOC"),
        name=clean_name,
        media_type=media_type,
        extension=extension,
        sha256=digest,
        byte_size=len(data),
        extracted_text=extracted_text,
        segments=segments,
        warnings=warnings,
        uploaded_at_utc=_utc_now(),
    )


def document_inventory_row(document: DocumentRecord) -> dict[str, Any]:
    return {
        "문서 ID": document.document_id,
        "파일명": document.name,
        "형식": document.extension or document.media_type,
        "크기(KB)": round(document.byte_size / 1024, 1),
        "근거 구간": len(document.segments),
        "추출 문자": len(document.extracted_text),
        "경고": len(document.warnings),
        "SHA-256": document.sha256[:16],
    }
