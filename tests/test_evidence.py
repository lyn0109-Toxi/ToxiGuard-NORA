from __future__ import annotations

import io
import unittest

from docx import Document
from openpyxl import Workbook
from reportlab.pdfgen import canvas

from nora.assertions import apply_reviewed_assertions, extract_assertions, reviewed_assertion_conflicts
from nora.evidence import extract_document
from nora.models import AssessmentInput


class EvidenceExtractionTests(unittest.TestCase):
    def test_text_document_and_assertions(self) -> None:
        text = """
        Product Name: GP-L-CT
        Model Name: Hepato Classifier
        Version: v3.0
        Endpoint: hepatotoxicity
        Sensitivity: 82%
        Specificity: 76%
        Prediction result: negative
        Training dataset mainly included small-molecule compounds.
        Primary human hepatocyte 2D cell assay result was negative.
        Positive control passed and negative control passed.
        """.encode("utf-8")
        record = extract_document(text, "evidence.txt", "text/plain")
        self.assertTrue(record.extracted_text)
        assertions = extract_assertions(record)
        paths = {item.field_path for item in assertions}
        self.assertIn("product.product_name", paths)
        self.assertIn("ai_model.model_name", paths)
        self.assertIn("ai_model.model_version", paths)
        self.assertIn("ai_model.sensitivity_percent", paths)
        self.assertIn("nam_assay.nam_type", paths)

    def test_docx_pdf_and_xlsx(self) -> None:
        doc = Document()
        doc.add_paragraph("Model Name: Human Hepato Model")
        doc.add_paragraph("Version: v1.2")
        buf = io.BytesIO()
        doc.save(buf)
        docx_record = extract_document(buf.getvalue(), "model.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        self.assertGreater(len(docx_record.segments), 0)

        pdf_buf = io.BytesIO()
        c = canvas.Canvas(pdf_buf)
        c.drawString(72, 760, "Model Name: PDF Hepato Model")
        c.drawString(72, 740, "Version: v2.0")
        c.save()
        pdf_record = extract_document(pdf_buf.getvalue(), "model.pdf", "application/pdf")
        self.assertGreater(len(pdf_record.segments), 0)

        wb = Workbook()
        ws = wb.active
        ws["A1"] = "Sensitivity"
        ws["B1"] = 90
        xlsx_buf = io.BytesIO()
        wb.save(xlsx_buf)
        xlsx_record = extract_document(xlsx_buf.getvalue(), "model.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        self.assertGreater(len(xlsx_record.segments), 0)

    def test_only_reviewed_assertions_are_applied(self) -> None:
        record = extract_document(b"Model Name: Reviewed Model\nVersion: v9.1", "model.txt", "text/plain")
        assertions = extract_assertions(record)
        for item in assertions:
            if item.field_path in {"ai_model.model_name", "ai_model.model_version"}:
                item.review_status = "승인"
        updated = apply_reviewed_assertions(AssessmentInput(), assertions)
        self.assertEqual(updated.ai_model.model_name, "Reviewed Model")
        self.assertEqual(updated.ai_model.model_version, "v9.1")
        self.assertTrue(updated.supporting_evidence.assertions_reviewed)

    def test_negated_evidence_does_not_create_positive_assertions(self) -> None:
        text = b"""
        Quantitative biodistribution was not available.
        Training data did not include siRNA nanomedicines.
        Free concentration and intracellular concentration were not measured.
        Kupffer cells and stellate cells were not included.
        Carrier-only control was not included.
        QIVIVE or PBPK translation was not performed.
        External validation was partially performed.
        """
        record = extract_document(text, "negated.txt", "text/plain")
        assertions = extract_assertions(record)
        pairs = {(item.field_path, item.proposed_value) for item in assertions}
        self.assertNotIn(("product.distribution_status", "정량적 자료"), pairs)
        self.assertNotIn(("supporting_evidence.quantitative_biodistribution", "true"), pairs)
        self.assertNotIn(("ai_model.domain_modalities", "나노의약품"), pairs)
        self.assertNotIn(("nam_assay.measured_exposure", "측정됨"), pairs)
        self.assertNotIn(("nam_assay.cell_types", "Kupffer cell"), pairs)
        self.assertNotIn(("nam_assay.cell_types", "Stellate cell"), pairs)
        self.assertNotIn(("nam_assay.carrier_only_control", "포함"), pairs)
        self.assertNotIn(("nam_assay.qivive_pbpk", "수행됨"), pairs)
        self.assertIn(("ai_model.external_validation", "부분적으로 확인"), pairs)

    def test_conflicting_reviewed_scalar_assertions_do_not_silently_overwrite(self) -> None:
        record = extract_document(
            b"Prediction result: negative. Prediction result: positive.",
            "conflict.txt",
            "text/plain",
        )
        assertions = [item for item in extract_assertions(record) if item.field_path == "ai_model.result"]
        for item in assertions:
            item.review_status = "승인"
        conflicts = reviewed_assertion_conflicts(assertions)
        self.assertIn("ai_model.result", conflicts)
        baseline = AssessmentInput()
        original = baseline.ai_model.result
        updated = apply_reviewed_assertions(baseline, assertions)
        self.assertEqual(updated.ai_model.result, original)


if __name__ == "__main__":
    unittest.main()
